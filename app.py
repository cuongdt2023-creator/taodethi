import streamlit as st
import io
import random
import re
from docx import Document
from docxcompose.composer import Composer
from copy import deepcopy

# ==================== CẤU HÌNH ====================
st.set_page_config(page_title="Trộn Đề Word: Bảo Toàn Tuyệt Đối", page_icon="💎", layout="wide")

st.markdown("""
<style>
    .main-header { text-align: center; color: #b91c1c; font-weight: bold; }
    .info-box { background-color: #fef2f2; border: 1px solid #fecaca; padding: 10px; border-radius: 5px; }
</style>
""", unsafe_allow_html=True)

# ==================== LOGIC XỬ LÝ WORD (THUẬT TOÁN TRIM) ====================

def get_difficulty(text):
    t = text.upper()
    if "#VDC" in t: return "VDC"
    if "#VD" in t: return "VD"
    if "#TH" in t: return "TH"
    if "#NB" in t: return "NB"
    return "NB"

def analyze_document_structure(file_bytes):
    """
    Quét file để tìm tọa độ (index) của các câu hỏi và phần đáp án.
    Không chỉnh sửa file ở bước này.
    """
    doc = Document(io.BytesIO(file_bytes))
    map_data = {
        "questions": [], # List of dict: {start_idx, end_idx, diff, part}
        "footer_start": -1, # Vị trí bắt đầu phần đáp án/hướng dẫn
        "p1_idx": -1, "p2_idx": -1, "p3_idx": -1
    }
    
    current_part = "P1"
    q_start = -1
    
    total_paras = len(doc.paragraphs)
    
    for i, p in enumerate(doc.paragraphs):
        txt = p.text.strip().upper()
        
        # 1. Nhận diện phần
        if txt.startswith("PHẦN 1") or txt.startswith("PHẦN I"): current_part = "P1"
        elif txt.startswith("PHẦN 2") or txt.startswith("PHẦN II"): current_part = "P2"
        elif txt.startswith("PHẦN 3") or txt.startswith("PHẦN III"): current_part = "P3"
        
        # 2. Nhận diện Bảng đáp án / Hướng dẫn (thường ở cuối)
        # Nếu gặp từ khóa này, coi như hết phần câu hỏi
        if "BẢNG ĐÁP ÁN" in txt or "HƯỚNG DẪN GIẢI" in txt or "LỜI GIẢI" in txt:
            if q_start != -1:
                 # Lưu câu hỏi cuối cùng trước khi vào phần đáp án
                diff = get_difficulty(doc.paragraphs[q_start].text)
                map_data["questions"].append({"range": (q_start, i), "diff": diff, "part": prev_part})
                q_start = -1
            map_data["footer_start"] = i
            break 

        # 3. Nhận diện Câu hỏi
        if re.match(r'^Câu\s*\d+', p.text, re.IGNORECASE):
            if q_start != -1:
                # Lưu câu hỏi trước đó
                diff = get_difficulty(doc.paragraphs[q_start].text)
                map_data["questions"].append({"range": (q_start, i), "diff": diff, "part": prev_part})
            
            q_start = i
            prev_part = current_part
            
    # Lưu câu hỏi cuối cùng nếu chưa gặp footer
    if q_start != -1 and map_data["footer_start"] == -1:
        diff = get_difficulty(doc.paragraphs[q_start].text)
        map_data["questions"].append({"range": (q_start, total_paras), "diff": diff, "part": prev_part})
    
    return map_data

def extract_content_by_trimming(file_bytes, keep_ranges):
    """
    Cốt lõi của phương pháp Triệt Để:
    Load file gốc -> Xóa TẤT CẢ các đoạn KHÔNG nằm trong keep_ranges -> Trả về Doc.
    keep_ranges: List các tuple (start, end) cần giữ lại.
    """
    doc = Document(io.BytesIO(file_bytes))
    
    # Tạo danh sách các index cần xóa (ngược lại với cần giữ)
    # Tư duy: Giữ lại những dòng user chọn, còn lại xóa hết.
    
    total_paras = len(doc.paragraphs)
    indices_to_keep = set()
    for start, end in keep_ranges:
        for i in range(start, end):
            indices_to_keep.add(i)
            
    # Xóa từ dưới lên trên để không làm lệch index
    for i in range(total_paras - 1, -1, -1):
        if i not in indices_to_keep:
            p = doc.paragraphs[i]
            # Xóa paragraph khỏi XML
            p._element.getparent().remove(p._element)
            
    return doc

# ==================== GIAO DIỆN CHÍNH ====================

st.markdown("<h1 class='main-header'>💎 Tạo Đề Chuẩn (Giữ Đáp Án & MathType)</h1>", unsafe_allow_html=True)
st.markdown("<div class='info-box'>⚠️ <b>Lưu ý quan trọng:</b> Chương trình sẽ tự động tìm phần <b>'BẢNG ĐÁP ÁN'</b> hoặc <b>'HƯỚNG DẪN GIẢI'</b> ở cuối mỗi file để gộp vào đề tổng hợp. Hãy đảm bảo file gốc có các mục này nếu bạn muốn giữ lại đáp án.</div>", unsafe_allow_html=True)

files = st.file_uploader("Bước 1: Tải các file chủ đề", type="docx", accept_multiple_files=True)

if files:
    # Phân tích cấu trúc (Metadata)
    if 'structs' not in st.session_state or len(st.session_state.structs) != len(files):
        with st.spinner("Đang quét cấu trúc file..."):
            st.session_state.structs = {}
            for f in files:
                f_bytes = f.read()
                st.session_state.structs[f.name] = {
                    "bytes": f_bytes,
                    "meta": analyze_document_structure(f_bytes)
                }

    st.subheader("Bước 2: Cấu hình số câu")
    configs = {}
    cols = st.columns(len(files))
    
    for i, fname in enumerate(st.session_state.structs.keys()):
        meta = st.session_state.structs[fname]["meta"]
        qs = meta["questions"]
        counts = {"P1": 0, "P2": 0, "P3": 0}
        for q in qs: counts[q["part"]] += 1
        
        has_footer = "✅ Có Đáp án" if meta["footer_start"] != -1 else "⚠️ Không thấy Đáp án"
        
        with cols[i]:
            st.info(f"📄 {fname[:15]}\n\n({has_footer})")
            configs[fname] = {
                "P1": st.number_input(f"P1 (Max {counts['P1']})", 0, 50, 0, key=f"p1_{fname}"),
                "P2": st.number_input(f"P2 (Max {counts['P2']})", 0, 50, 0, key=f"p2_{fname}"),
                "P3": st.number_input(f"P3 (Max {counts['P3']})", 0, 50, 0, key=f"p3_{fname}")
            }

    if st.button("🚀 XUẤT ĐỀ THI HOÀN CHỈNH", type="primary", use_container_width=True):
        status = st.empty()
        try:
            # 1. Tạo Master Doc từ file đầu tiên (Xóa sạch nội dung, giữ định dạng)
            base_bytes = list(st.session_state.structs.values())[0]["bytes"]
            master_doc = Document(io.BytesIO(base_bytes))
            for p in master_doc.paragraphs: p._element.getparent().remove(p._element)
            
            composer = Composer(master_doc)
            
            titles = {
                "P1": "PHẦN I. Câu trắc nghiệm nhiều phương án lựa chọn.",
                "P2": "PHẦN II. Câu trắc nghiệm đúng sai.",
                "P3": "PHẦN III. Câu trắc nghiệm trả lời ngắn."
            }
            
            # --- XỬ LÝ GỘP CÂU HỎI ---
            global_q_idx = {"P1": 1, "P2": 1, "P3": 1}
            
            for p_key in ["P1", "P2", "P3"]:
                status.write(f"⏳ Đang xử lý {titles[p_key]}...")
                
                # Gom tất cả request cho phần này
                part_requests = [] # List of {fname, q_data}
                
                for fname, cfg in configs.items():
                    data = st.session_state.structs[fname]
                    pool = [q for q in data["meta"]["questions"] if q["part"] == p_key]
                    num = min(cfg[p_key], len(pool))
                    if num > 0:
                        selected = random.sample(pool, num)
                        for q in selected:
                            part_requests.append({"fname": fname, "q": q, "bytes": data["bytes"]})
                
                if part_requests:
                    # Thêm tiêu đề phần vào Master
                    master_doc.add_paragraph(titles[p_key]).bold = True
                    random.shuffle(part_requests)
                    
                    for req in part_requests:
                        # TRICK: Mở file gốc -> Xóa hết trừ câu hỏi này -> Append vào Master
                        # Cách này giữ 100% MathType/Ảnh của câu hỏi đó
                        q_range = req["q"]["range"]
                        q_doc = extract_content_by_trimming(req["bytes"], [q_range])
                        
                        # Đánh lại số câu
                        # Vì q_doc đã bị trim, câu hỏi chắc chắn nằm ở đoạn đầu
                        for p in q_doc.paragraphs:
                            if re.match(r'^Câu\s*\d+', p.text, re.IGNORECASE):
                                p.text = re.sub(r'^Câu\s*\d+', f"Câu {global_q_idx[p_key]}", p.text, flags=re.IGNORECASE)
                                p.text = re.sub(r'#(NB|TH|VD|VDC)', '', p.text)
                                break
                        
                        global_q_idx[p_key] += 1
                        composer.append(q_doc)

            # --- XỬ LÝ GỘP ĐÁP ÁN (FOOTER) ---
            status.write("⏳ Đang tổng hợp Đáp án & Hướng dẫn giải...")
            master_doc.add_page_break()
            master_doc.add_paragraph("--- TỔNG HỢP ĐÁP ÁN & HƯỚNG DẪN ---").bold = True
            
            for fname in configs.keys():
                data = st.session_state.structs[fname]
                footer_start = data["meta"]["footer_start"]
                
                # Nếu file này có phần đáp án và chúng ta có lấy câu hỏi từ file này
                total_picked = sum(configs[fname].values())
                if footer_start != -1 and total_picked > 0:
                    master_doc.add_paragraph(f"Nguồn: {fname}").italic = True
                    
                    # Cắt lấy phần đuôi từ footer_start đến hết
                    total_len = len(Document(io.BytesIO(data["bytes"])).paragraphs)
                    footer_doc = extract_content_by_trimming(data["bytes"], [(footer_start, total_len)])
                    composer.append(footer_doc)
                    master_doc.add_paragraph("-" * 20)

            # Xuất file
            out = io.BytesIO()
            master_doc.save(out)
            
            status.empty()
            st.success("✅ Xử lý hoàn tất! Cấu trúc, Công thức và Đáp án đã được bảo toàn.")
            st.download_button("📥 Tải đề thi (.docx)", out.getvalue(), "De_Thi_Tiet_Kiem_Format.docx")
            
        except Exception as e:
            st.error(f"Lỗi hệ thống: {str(e)}")
